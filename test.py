from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Tạo document
doc = Document()

# Cài đặt margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(2.5)
    section.bottom_margin = Inches(2.5)
    section.left_margin = Inches(2.5)
    section.right_margin = Inches(2.5)

# Style tổng quát
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# ============ CHƯƠNG 2 ============
heading = doc.add_heading('Chương 2: Cơ sở lý thuyết', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

# 2.1
heading21 = doc.add_heading('2.1. Tổng quan về phương pháp lọc cộng tác', level=2)

p1 = doc.add_paragraph(
    'Lọc cộng tác (Collaborative Filtering – CF) là một trong những kỹ thuật quan trọng nhất trong các hệ thống gợi ý hiện đại. Phương pháp này được sử dụng để dự đoán sở thích của người dùng đối với các sản phẩm hoặc nội dung mà họ chưa từng tương tác. CF khai thác thông tin từ dữ liệu hành vi của nhiều người dùng trong quá khứ để đưa ra các gợi ý phù hợp cho từng cá nhân.'
)
p1.paragraph_format.line_spacing = 1.5

p2 = doc.add_paragraph(
    'Nguyên lý cơ bản của CF dựa trên giả định rằng những người dùng có hành vi tương tự trong quá khứ sẽ có xu hướng đưa ra những lựa chọn tương tự trong tương lai. Cụ thể, trong bối cảnh hệ thống gợi ý phim (như đã nêu ở Chương 1), nếu hai người dùng đã đánh giá một số bộ phim giống nhau và có mức độ yêu thích tương tự, thì họ cũng có khả năng cao sẽ thích những bộ phim khác mà họ chưa xem. Do đó, thay vì phân tích nội dung của sản phẩm (như thể loại, diễn viên), CF tập trung vào việc phân tích mối quan hệ giữa người dùng và sản phẩm thông qua dữ liệu tương tác như rating, lượt xem hoặc hành vi sử dụng.'
)
p2.paragraph_format.line_spacing = 1.5

p3 = doc.add_paragraph(
    'Một đặc điểm nổi bật của CF là không yêu cầu thông tin mô tả chi tiết về sản phẩm, giúp phương pháp này có thể áp dụng linh hoạt trong nhiều lĩnh vực khác nhau (phim, nhạc, sách, sản phẩm e-commerce, v.v.). Tuy nhiên, do phụ thuộc hoàn toàn vào dữ liệu tương tác, CF thường gặp phải các vấn đề như dữ liệu thưa (sparsity) – như đã đề cập trong Chương 1 – và khó khăn trong việc xử lý các trường hợp người dùng hoặc sản phẩm mới (cold-start problem).'
)
p3.paragraph_format.line_spacing = 1.5

p4 = doc.add_paragraph()
p4.add_run('Trong các hệ thống thực tế, CF thường được triển khai theo hai hướng chính:\n').bold = False
p4_text = p4.add_run(
    '1) Phương pháp dựa trên lân cận (neighborhood-based): sử dụng thông tin từ các neighbors (người dùng hoặc sản phẩm tương tự) để dự đoán.\n'
    '2) Phương pháp dựa trên mô hình (model-based): sử dụng các mô hình như matrix factorization, neural networks để tìm các latent factors ẩn trong dữ liệu.\n\n'
)
p4_text.italic = False
p4.paragraph_format.line_spacing = 1.5

p5 = doc.add_paragraph(
    'Trong phạm vi của đề tài này (như định hướng ở phần 1.3), phương pháp dựa trên lân cận được lựa chọn do tính đơn giản, dễ triển khai và phù hợp với dữ liệu rating tường minh từ bộ dữ liệu MovieLens. Cụ thể, hai biến thể chính được áp dụng là User-based CF và Item-based CF, sẽ được trình bày chi tiết ở các phần 2.4 và 2.5.'
)
p5.paragraph_format.line_spacing = 1.5

# 2.2
heading22 = doc.add_heading('2.2. Biểu diễn dữ liệu trong lọc cộng tác', level=2)

p6 = doc.add_paragraph(
    'Trong CF, dữ liệu thường được biểu diễn dưới dạng ma trận hai chiều gọi là ma trận user–item (utility matrix hoặc rating matrix). Mỗi hàng của ma trận tương ứng với một người dùng, mỗi cột tương ứng với một sản phẩm (phim), và giá trị tại mỗi ô biểu thị mức độ quan tâm của người dùng đối với sản phẩm đó, thường dưới dạng một con số rating (ví dụ: từ 1 đến 5 sao).'
)
p6.paragraph_format.line_spacing = 1.5

p7 = doc.add_paragraph(
    'Giả sử có m người dùng và n sản phẩm, khi đó ma trận user–item có kích thước m × n. Ví dụ, với bộ dữ liệu MovieLens được sử dụng trong đề tài này, có thể có hàng triệu người dùng và hàng chục nghìn bộ phim, tạo thành một ma trận cực kỳ lớn.'
)
p7.paragraph_format.line_spacing = 1.5

p8 = doc.add_paragraph(
    'Tuy nhiên, trong thực tế, mỗi người dùng chỉ tương tác (đánh giá) với một số ít sản phẩm, dẫn đến việc phần lớn các phần tử trong ma trận không có giá trị (được ký hiệu là "?"). Hiện tượng này được gọi là tính thưa (sparsity) của dữ liệu. Mức độ thưa có thể rất cao – chẳng hạn, nếu mỗi người dùng chỉ đánh giá 50 bộ phim từ 10.000 bộ phim sẵn có, thì độ thưa sẽ là 99.5%.'
)
p8.paragraph_format.line_spacing = 1.5

p9 = doc.add_paragraph(
    'Tính thưa của ma trận gây ra nhiều khó khăn trong quá trình tính toán, đặc biệt là khi xác định độ tương đồng giữa các người dùng hoặc sản phẩm. Để giải quyết vấn đề này, dữ liệu thường được lưu trữ dưới dạng sparse matrix (ma trận thưa), trong đó chỉ lưu các giá trị đã biết cùng với vị trí của chúng. Cách biểu diễn này giúp giảm đáng kể bộ nhớ sử dụng và tăng hiệu quả tính toán.'
)
p9.paragraph_format.line_spacing = 1.5

p10 = doc.add_paragraph(
    'Ngoài ra, trong nhiều trường hợp, dữ liệu rating được chuẩn hóa (normalization) theo từng người dùng nhằm loại bỏ sự khác biệt trong thang đánh giá. Ví dụ, một số người dùng có xu hướng đánh giá cao hơn trung bình (optimistic raters), trong khi những người khác lại có xu hướng đánh giá thấp hơn (pessimistic raters). Việc chuẩn hóa giúp đưa các dữ liệu này về cùng một hệ quy chiếu, từ đó cải thiện độ chính xác của các phép tính tiếp theo.'
)
p10.paragraph_format.line_spacing = 1.5

# 2.3
heading23 = doc.add_heading('2.3. Độ tương đồng trong lọc cộng tác', level=2)

p11 = doc.add_paragraph(
    'Việc xác định độ tương đồng (similarity) giữa các phần tử là bước cốt lõi trong phương pháp CF dựa trên lân cận. Độ tương đồng cho phép hệ thống xác định các đối tượng gần nhau, từ đó sử dụng thông tin của chúng để dự đoán giá trị chưa biết.'
)
p11.paragraph_format.line_spacing = 1.5

p12 = doc.add_paragraph(
    'Trong bối cảnh CF, độ tương đồng có thể được tính giữa hai người dùng hoặc giữa hai sản phẩm:'
)
p12.paragraph_format.line_spacing = 1.5

bullet1 = doc.add_paragraph('Khi tính độ tương đồng giữa hai người dùng, mỗi người dùng được biểu diễn dưới dạng một vector rating trên các sản phẩm (một hàng của ma trận user–item).', style='List Bullet')
bullet1.paragraph_format.line_spacing = 1.5

bullet2 = doc.add_paragraph('Khi tính độ tương đồng giữa hai sản phẩm, mỗi sản phẩm được biểu diễn dưới dạng một vector rating từ các người dùng (một cột của ma trận user–item).', style='List Bullet')
bullet2.paragraph_format.line_spacing = 1.5

p13 = doc.add_paragraph(
    'Một trong những độ đo phổ biến nhất là cosine similarity. Phương pháp này đo độ giống nhau giữa hai vector dựa trên góc giữa chúng trong không gian đa chiều. Nếu hai vector có hướng giống nhau, giá trị cosine sẽ gần 1; nếu vuông góc, giá trị sẽ gần 0; và nếu ngược hướng, giá trị sẽ gần -1.'
)
p13.paragraph_format.line_spacing = 1.5

p14 = doc.add_paragraph(
    'Công thức tính cosine similarity giữa hai vector u và v là:'
)
p14.paragraph_format.line_spacing = 1.5

# Thêm công thức
p_formula = doc.add_paragraph()
p_formula.add_run('sim(u, v) = cos(u, v) = (u').font.size = Pt(11)
p_formula.add_run('T').font.size = Pt(9)
p_formula.add_run(' · v) / (||u||').font.size = Pt(11)
p_formula.add_run('2').font.size = Pt(9)
p_formula.add_run(' · ||v||').font.size = Pt(11)
p_formula.add_run('2').font.size = Pt(9)
p_formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

p15 = doc.add_paragraph(
    'Độ tương đồng của hai phần tử nằm trong khoảng [-1, 1], trong đó giá trị càng lớn thì độ tương đồng càng cao.'
)
p15.paragraph_format.line_spacing = 1.5

p16 = doc.add_paragraph(
    'Trong thực tế, để tăng độ tin cậy của độ tương đồng, người ta thường chỉ xét các cặp có đủ số lượng phần tử chung tối thiểu (common items hoặc common users). Ngoài ra, một số kỹ thuật như shrinkage cũng được áp dụng nhằm giảm ảnh hưởng của các trường hợp có ít dữ liệu. Việc lựa chọn độ đo tương đồng phù hợp có ảnh hưởng lớn đến chất lượng của hệ thống gợi ý.'
)
p16.paragraph_format.line_spacing = 1.5

# 2.4
heading24 = doc.add_heading('2.4. Phương pháp User-based Collaborative Filtering', level=2)

p17 = doc.add_paragraph(
    'User-based CF là phương pháp dựa trên việc tìm kiếm các người dùng có hành vi tương tự với người dùng mục tiêu. Ý tưởng chính của phương pháp này là: nếu người dùng u1 và u2 có độ tương đồng cao (tức là họ đã đánh giá các bộ phim giống nhau một cách tương tự), thì việc dự đoán rating của u1 cho một bộ phim chưa xem có thể dựa trên rating của u2 cho bộ phim đó. Như vậy, hệ thống sử dụng thông tin từ các người dùng "gần" để dự đoán sở thích của người dùng cần gợi ý.'
)
p17.paragraph_format.line_spacing = 1.5

p18 = doc.add_paragraph('Quy trình thực hiện phương pháp User-based CF bao gồm các bước chính:')
p18.paragraph_format.line_spacing = 1.5

list_items = [
    'Xây dựng ma trận user–item từ dữ liệu rating có sẵn.',
    'Tính toán độ tương đồng (sử dụng cosine similarity) giữa người dùng mục tiêu u và tất cả các người dùng khác.',
    'Lựa chọn tập K người dùng gần nhất (K-nearest neighbors, KNN) dựa trên độ tương đồng cao nhất.',
    'Dự đoán rating dựa trên thông tin từ K người dùng này.'
]

for i, item in enumerate(list_items, 1):
    p = doc.add_paragraph(f'{i}. {item}', style='List Number')
    p.paragraph_format.line_spacing = 1.5

p19 = doc.add_paragraph(
    'Phương pháp này có mối liên hệ chặt chẽ với thuật toán K-Nearest Neighbors (KNN), trong đó việc lựa chọn K neighbors đóng vai trò quan trọng trong việc quyết định chất lượng dự đoán. Nếu K quá nhỏ, dự đoán có thể bị thiên lệch; nếu K quá lớn, có thể bao gồm các neighbors không liên quan.'
)
p19.paragraph_format.line_spacing = 1.5

p20 = doc.add_paragraph(
    'Giá trị dự đoán rating được tính bằng trung bình có trọng số của các rating từ các người dùng tương tự, trong đó trọng số là độ tương đồng. Công thức được sử dụng là:'
)
p20.paragraph_format.line_spacing = 1.5

p_formula2 = doc.add_paragraph(
    'r̂(u,i) = r̄u + Σ[sim(u, u′) · (r(u′,i) − r̄u′)] / Σ|sim(u, u′)|'
)
p_formula2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_formula2.paragraph_format.line_spacing = 1.5

advantages = doc.add_paragraph('Ưu điểm của phương pháp này:')
advantages.paragraph_format.line_spacing = 1.5

adv_items = [
    'Dễ hiểu và dễ triển khai, phù hợp với bài toán nhập môn trí tuệ nhân tạo.',
    'Mang lại mức độ cá nhân hóa cao, vì sử dụng trực tiếp hành vi của từng người dùng.',
    'Có thể khám phá ra những item bất ngờ mà người dùng chưa biết thông qua các neighbors.'
]

for adv in adv_items:
    p = doc.add_paragraph(adv, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5

disadvantages = doc.add_paragraph('Hạn chế của phương pháp:')
disadvantages.paragraph_format.line_spacing = 1.5

disadv_items = [
    'Khi số lượng người dùng lớn, việc tính toán độ tương đồng giữa tất cả các cặp user có thể trở nên tốn kém (độ phức tạp O(m²n)).',
    'Gặp vấn đề cold-start với người dùng mới (người dùng chưa đánh giá bất kỳ item nào).',
    'Dữ liệu thưa làm giảm độ chính xác của việc tính toán độ tương đồng.'
]

for disadv in disadv_items:
    p = doc.add_paragraph(disadv, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5

# 2.5
heading25 = doc.add_heading('2.5. Phương pháp Item-based Collaborative Filtering', level=2)

p21 = doc.add_paragraph(
    'Khác với User-based CF, phương pháp Item-based CF tập trung vào việc xác định độ tương đồng giữa các sản phẩm (phim). Trong phương pháp này, mỗi sản phẩm được biểu diễn như một vector dựa trên các rating từ người dùng.'
)
p21.paragraph_format.line_spacing = 1.5

p22 = doc.add_paragraph(
    'Ý tưởng chính của phương pháp này là: nếu một người dùng đã thích một bộ phim (đánh giá cao), thì họ có khả năng sẽ thích các bộ phim tương tự với bộ phim đó. Ví dụ, nếu người dùng A yêu thích phim "Avatar", thì hệ thống có thể gợi ý những phim sci-fi có đánh giá tương tự "Avatar" từ các user khác.'
)
p22.paragraph_format.line_spacing = 1.5

p23 = doc.add_paragraph('Quy trình thực hiện Item-based CF bao gồm:')
p23.paragraph_format.line_spacing = 1.5

item_list = [
    'Tính toán độ tương đồng giữa các sản phẩm (phim) sử dụng cosine similarity.',
    'Xác định các sản phẩm gần nhất (most similar items) với các sản phẩm mà người dùng đã đánh giá.',
    'Dự đoán rating dựa trên các sản phẩm mà người dùng đã đánh giá.'
]

for i, item in enumerate(item_list, 1):
    p = doc.add_paragraph(f'{i}. {item}', style='List Number')
    p.paragraph_format.line_spacing = 1.5

p24 = doc.add_paragraph(
    'So với User-based CF, phương pháp này có một số ưu điểm đáng kể:'
)
p24.paragraph_format.line_spacing = 1.5

adv_item_based = [
    'Ổn định hơn: Số lượng sản phẩm thường ít thay đổi hơn so với số lượng người dùng. Do đó, việc tính toán và lưu trữ ma trận tương đồng giữa các sản phẩm trở nên ổn định hơn theo thời gian.',
    'Tốc độ xử lý nhanh: Ma trận tương đồng giữa các sản phẩm có thể được tính trước (offline), giúp tăng tốc độ phản hồi khi hệ thống hoạt động (online). Đây là lợi thế lớn cho các ứng dụng thực tế yêu cầu độ trễ thấp.',
    'Hiệu quả với số lượng sản phẩm ít: Độ phức tạp là O(n²m), thường tốt hơn User-based CF khi n < m.'
]

for i, adv in enumerate(adv_item_based, 1):
    p = doc.add_paragraph(f'{i}. {adv}', style='List Number')
    p.paragraph_format.line_spacing = 1.5

disadv_item_based = doc.add_paragraph('Hạn chế của phương pháp:')
disadv_item_based.paragraph_format.line_spacing = 1.5

disadv_items_ib = [
    'Không tận dụng trực tiếp mối quan hệ giữa các người dùng, do đó có thể bỏ lỡ các gợi ý bất ngờ.',
    'Gặp vấn đề cold-start với sản phẩm mới (sản phẩm chưa được ai đánh giá).',
    'Có khuynh hướng gợi ý các sản phẩm tương tự (filter bubble), làm giảm độ đa dạng của gợi ý.'
]

for disadv in disadv_items_ib:
    p = doc.add_paragraph(disadv, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5

# 2.6
heading26 = doc.add_heading('2.6. Vấn đề Cold-start và Độ đo Đánh giá', level=2)

p25 = doc.add_paragraph(
    'Cold-start problem là một trong những thách thức quan trọng của CF. Vấn đề này xảy ra khi: (1) người dùng mới tham gia nhưng chưa đánh giá bất kỳ sản phẩm nào, (2) sản phẩm mới được thêm vào nhưng chưa được ai đánh giá, hoặc (3) dữ liệu quá thưa khiến tính toán độ tương đồng không chính xác.'
)
p25.paragraph_format.line_spacing = 1.5

p26 = doc.add_paragraph(
    'Để giảm nhẹ vấn đề cold-start, các giải pháp phổ biến bao gồm: hybrid approach (kết hợp CF với content-based), demographic filtering (sử dụng thông tin nhân khẩu học), pop-based fallback (gợi ý sản phẩm phổ biến), và knowledge-based recommendation (yêu cầu user cung cấp thông tin). Trong đề tài này, giải pháp dự phòng là gợi ý những phim có điểm đánh giá trung bình cao.'
)
p26.paragraph_format.line_spacing = 1.5

p27 = doc.add_paragraph('Để đánh giá chất lượng hệ thống gợi ý, các chỉ số chính bao gồm:')
p27.paragraph_format.line_spacing = 1.5

metrics = [
    'RMSE (Root Mean Squared Error): Đo sai số dự đoán rating, giá trị càng nhỏ càng tốt.',
    'Precision@K: Tỷ lệ sản phẩm được gợi ý thực sự được user quan tâm trong top-K.',
    'Recall@K: Tỷ lệ sản phẩm user quan tâm được tìm thấy trong top-K.'
]

for metric in metrics:
    p = doc.add_paragraph(metric, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5

# Kết luận
p28 = doc.add_paragraph()
p28.add_run('Kết luận\n').bold = True
p28_text = p28.add_run(
    'Chương này đã trình bày các kiến thức nền tảng về Collaborative Filtering, bao gồm nguyên lý hoạt động, cách biểu diễn dữ liệu, các phương pháp User-based CF và Item-based CF, cùng với các vấn đề và độ đo đánh giá liên quan. Các kiến thức này sẽ là cơ sở cho việc triển khai hệ thống gợi ý phim trong Chương 3.'
)
p28.paragraph_format.line_spacing = 1.5

# Lưu file
doc.save('Chuong2_CoSoLyThuyetCaiBien.docx')
print("✅ File Word tạo thành công: Chuong2_CoSoLyThuyetCaiBien.docx")